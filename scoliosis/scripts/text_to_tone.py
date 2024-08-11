import argparse
import os
import nltk
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.corpus import stopwords
from nltk.stem import PorterStemmer, WordNetLemmatizer
from nltk import pos_tag, ne_chunk
from textblob import TextBlob
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import seaborn as sns
from collections import Counter
from docx import Document
from transformers import pipeline
import shap

# Ensure the necessary NLTK data is downloaded
nltk.download('punkt')
nltk.download('stopwords')
nltk.download('averaged_perceptron_tagger')
nltk.download('maxent_ne_chunker')
nltk.download('words')

# Default text if no input is provided
DEFAULT_TEXT = "scoliosis/scripts/docs/20240616_Trial1.docx"

# Custom stop words
CUSTOM_STOP_WORDS = {
    'yeah', 'oh', 'really', 'okay', 'know', 'gotcha', 'like', 'kind', 'right', 'feel', 'think', 'honestly', 
    'anything', 'want', "n't", "'s", "'m", "'ve", "'ll", "'d", 'ca', 'wo', 're', 'nt', 'na', 'gon'
}

def read_default_text():
    """
    Reads the default text from the specified DOCX file.

    This function reads the text from the DOCX file specified in the DEFAULT_TEXT variable.
    It returns the full text as a single string. If there is an error reading the file,
    it prints an error message and returns an empty string.

    Returns:
        str: The full text from the DOCX file or an empty string if an error occurs.
    """
    try:
        print(f"Attempting to read default text from: {DEFAULT_TEXT}")

        if not os.path.exists(DEFAULT_TEXT):
            raise FileNotFoundError(f"Default text file not found: {DEFAULT_TEXT}")

        doc = Document(DEFAULT_TEXT)
        print(doc)
        
        print(f"Successfully read the document: {DEFAULT_TEXT}")
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        return ' '.join(full_text)
    except Exception as e:
        print(f"Error reading default text: {e}")
        return ""

def remove_stopwords(text):
    """
    Removes stop words from the given text.

    Args:
        text (str): The input text.

    Returns:
        str: The text with stop words removed.
    """
    words = word_tokenize(text)
    stop_words = set(stopwords.words('english')).union(CUSTOM_STOP_WORDS)
    filtered_words = [word for word in words if word.lower() not in stop_words]
    return ' '.join(filtered_words)

def clean_tokens(tokens):
    """
    Cleans tokens by removing punctuation and non-alphabetic tokens.

    Args:
        tokens (list): The list of tokens.

    Returns:
        list: The cleaned list of tokens.
    """
    tokens = [word for word in tokens if word.isalpha()]
    return tokens

def tokenize(text):
    """
    Tokenizes the text into words and sentences, and visualizes the top 10 most common words.

    Args:
        text (str): The input text.
    """
    words = word_tokenize(text)
    words = clean_tokens(words)
    sentences = sent_tokenize(text)
    print("Words:", words)
    print("Sentences:", sentences)

    # Visualization
    word_freq = Counter(words)
    common_words = word_freq.most_common(10)
    words, counts = zip(*common_words)
    plt.figure(figsize=(10, 5))
    sns.barplot(x=list(counts), y=list(words))
    plt.title("Top 10 Most Common Words")
    plt.show()

def top_words(text):
    """
    Removes stop words and tokenizes the filtered text.

    Args:
        text (str): The input text.
    """
    filtered_text = remove_stopwords(text)
    return tokenize(filtered_text)

def stem(text):
    """
    Applies stemming to the text and prints the stemmed words.

    Args:
        text (str): The input text.
    """
    words = word_tokenize(text)
    words = clean_tokens(words)
    stemmer = PorterStemmer()
    stemmed_words = [stemmer.stem(word) for word in words]
    print("Stemmed Words:", stemmed_words)

def lemmatize(text):
    """
    Applies lemmatization to the text and prints the lemmatized words.

    Args:
        text (str): The input text.
    """
    words = word_tokenize(text)
    words = clean_tokens(words)
    lemmatizer = WordNetLemmatizer()
    lemmatized_words = [lemmatizer.lemmatize(word) for word in words]
    print("Lemmatized Words:", lemmatized_words)

def pos_tagging(text):
    """
    Performs part-of-speech tagging and visualizes the POS tags.

    Args:
        text (str): The input text.
    """
    words = word_tokenize(text)
    words = clean_tokens(words)
    pos_tags = pos_tag(words)
    print("POS Tags:", pos_tags)

    # Visualization
    pos_counts = Counter(tag for word, tag in pos_tags)
    tags, counts = zip(*pos_counts.items())
    plt.figure(figsize=(10, 5))
    sns.barplot(x=list(counts), y=list(tags))
    plt.title("Part of Speech Tagging")
    plt.show()

def named_entity_recognition(text):
    """
    Performs named entity recognition on the text.

    Args:
        text (str): The input text.
    """
    words = word_tokenize(text)
    words = clean_tokens(words)
    pos_tags = pos_tag(words)
    named_entities = ne_chunk(pos_tags)
    print("Named Entities:", named_entities)

def sentiment_analysis(text):
    """
    Analyzes the sentiment of the text using TextBlob and visualizes the results.

    Args:
        text (str): The input text.
    """
    filtered_text = remove_stopwords(text)
    blob = TextBlob(filtered_text)
    sentiment = blob.sentiment
    print("Sentiment:", sentiment)

    # Visualization
    plt.figure(figsize=(10, 5))
    plt.bar(
        ['Polarity', 'Subjectivity'], [sentiment.polarity, sentiment.subjectivity], color=['blue', 'orange']
    )
    plt.ylim([-1, 1])
    plt.title('Sentiment Analysis')
    plt.show()

def word_cloud(text):
    """
    Generates a word cloud from the text.

    Args:
        text (str): The input text.
    """
    filtered_text = remove_stopwords(text)
    words = word_tokenize(filtered_text)
    stemmer = PorterStemmer()
    stemmed_words = [stemmer.stem(word) for word in words]
    wordcloud = WordCloud(width=800, height=400, background_color='white').generate(' '.join(stemmed_words))
    plt.figure(figsize=(10, 5))
    plt.imshow(wordcloud, interpolation='bilinear')
    plt.axis('off')
    plt.show()

def bert_sentiment_analysis(texts):
    """
    Performs sentiment analysis using a BERT-based pipeline.

    Args:
        texts (list): The list of texts.

    Returns:
        list: The sentiment analysis results.
    """
    sentiment_pipeline = pipeline('sentiment-analysis')
    results = sentiment_pipeline(texts)
    return results

def explain_sentiment_analysis(texts):
    """
    Explains sentiment analysis using SHAP values.

    Args:
        texts (list): The list of texts.

    Returns:
        shap.Explanation: The SHAP values.
    """
    sentiment_pipeline = pipeline('sentiment-analysis')
    explainer = shap.Explainer(sentiment_pipeline)
    shap_values = explainer(texts)
    return shap_values

def visualize_bert_results(sentiment_scores, shap_values, texts):
    """
    Visualizes the sentiment scores and SHAP values for the texts.

    Args:
        sentiment_scores (list): The sentiment scores.
        shap_values (shap.Explanation): The SHAP values.
        texts (list): The list of texts.
    """
    # Plot sentiment scores
    plt.figure(figsize=(10, 5))
    plt.hist([score['score'] for score in sentiment_scores], bins=20, color='blue', alpha=0.7)
    plt.xlabel('Sentiment Score')
    plt.ylabel('Frequency')
    plt.title('Distribution of Sentiment Scores')
    plt.show()

    # Plot SHAP values
    for i in range(len(texts)):
        shap.plots.text(shap_values[i], display=True)
        plt.title(f"Text: {texts[i][:100]}...")  # Display first 100 characters of the text
        plt.show()

def process_bert(text):
    """
    Processes the text using BERT-based sentiment analysis and SHAP explanations.

    Args:
        text (str): The input text.
    """
    # Assuming text is a single string, split it into sentences or smaller chunks for analysis
    sentences = sent_tokenize(text)
    # Perform sentiment analysis
    sentiment_scores = bert_sentiment_analysis(sentences)

    # Explain sentiment analysis with SHAP
    shap_values = explain_sentiment_analysis(sentences)

    # Visualize results
    visualize_bert_results(sentiment_scores, shap_values, sentences)

# Main function to parse arguments and call the appropriate function
def main():
    """
    Main function to parse arguments and call the appropriate text processing function.
    """
    parser = argparse.ArgumentParser(description="Word Analyzer Tool")
    parser.add_argument("function", choices=["tokenize", "stem", "lemmatize", "pos_tagging", "ner", "sentiment", "wordcloud", "top", "bert"], help="Function to execute")
    parser.add_argument("--text", type=str, help="Text to analyze or path to a text file", default=None)
    args = parser.parse_args()

    # Read text from file if path is provided
    if args.text and args.text.endswith(".docx"):
        try:
            doc = Document(args.text)
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            text = ' '.join(full_text)
        except Exception as e:
            print(f"Error reading text file: {e}")
            text = read_default_text()
    elif args.text:
        text = args.text
    else:
        text = read_default_text()

    if args.function == "tokenize":
        tokenize(text)
    elif args.function == "stem":
        stem(text)
    elif args.function == "lemmatize":
        lemmatize(text)
    elif args.function == "pos_tagging":
        pos_tagging(text)
    elif args.function == "ner":
        named_entity_recognition(text)
    elif args.function == "sentiment":
        sentiment_analysis(text)
    elif args.function == "wordcloud":
        word_cloud(text)
    elif args.function == "top":
        top_words(text)
    elif args.function == "bert":
        process_bert(text)

if __name__ == "__main__":
    main()
